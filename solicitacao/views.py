import os
import tempfile
from datetime import datetime
from pprint import pprint

from decouple import config
from django.contrib.auth.forms import PasswordResetForm
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.models import User
from django.contrib.auth.tokens import default_token_generator
from django.core.mail import BadHeaderError, send_mail
from django.db.models import Q
from django.http import HttpRequest, HttpResponse
from django.shortcuts import redirect, render
from django.template.loader import get_template, render_to_string
from django.urls import reverse, reverse_lazy
from django.utils.encoding import force_bytes
from django.utils.http import urlsafe_base64_encode
from django.views.generic import (CreateView, DeleteView, DetailView,
                                  UpdateView, View)
from django.views.generic.list import ListView
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from htmldocx import HtmlToDocx
from secretaria.models import Secretaria
from solicita.storage_backends import PublicMediaStorage
from xhtml2pdf import pisa

from .forms import PasswordResetForm, SolicitacaoForm
from .models import Solicitacao


class SolicitacaoIndex(LoginRequiredMixin, ListView):
    model = Solicitacao
    template_name = 'solicitacao/index.html'
    context_object_name = 'solicitacoes'
    paginate_by = 10

    def get_queryset(self):
        query_set = super().get_queryset()
        query_set = query_set.select_related('secretaria')
        return query_set

    class Meta:
        ordering = ['-id']


class SolicitacaoDetailView(LoginRequiredMixin, DetailView):
    model = Solicitacao
    template_name = "solicitacao/solicitacao.html"
    context_object_name = 'solicitacao'


class SolicitacaoPDF(LoginRequiredMixin, View):
    model = Solicitacao
    template_name = 'solicitacao/render_pdf.html'
    context_object_name = 'solicitacao'

    def get(self, *args, **kwargs):
        template_path = self.template_name
        solicitacao = Solicitacao.objects.filter(pk=kwargs.get('pk'))
        solicitacao = solicitacao.select_related('secretaria').first()
        context = {'solicitacao': solicitacao}
        # Create a Django response object, and specify content_type as pdf
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="report.pdf"'
        # find the template and render it.
        template = get_template(template_path)
        html = template.render(context)
        # create a pdf
        pisa_status = pisa.CreatePDF(
            html, dest=response)
        # if error then show some funny view
        if pisa_status.err:
            return HttpResponse('We had some errors <pre>' + html + '</pre>')
        return response


class SolicitacaoDOCX(LoginRequiredMixin, View):
    model = Solicitacao
    template_name = 'solicitacao/render_doc.html'
    context_object_name = 'solicitacao'

    def get(self, *args, **kwargs):
        template_path = self.template_name
        solicitacao = Solicitacao.objects.filter(pk=kwargs.get('pk'))
        solicitacao = solicitacao.select_related('secretaria').first()
        context = {'solicitacao': solicitacao}
        template = get_template(template_path)
        html = template.render(context)
        media_storage: PublicMediaStorage = PublicMediaStorage()
        default_doc: PublicMediaStorage = media_storage.open('default.docx')
        document: Document = Document(docx=default_doc)
        docx: HtmlToDocx = HtmlToDocx()
        docx.add_html_to_document(html, document)
        # pprint(dir(solicitacao.first().secretaria.prefeitura.get_logo))
        # pprint(vars(solicitacao.first().secretaria.prefeitura.get_logo))
        section = document.sections[0]
        header = section.header
        # paragraph = header.paragraphs[0]
        # paragraph.add_picture(
        #     solicitacao.first().secretaria.prefeitura.get_logo)
        # paragraph.text = '\t{}'.format(solicitacao.secretaria.prefeitura)
        # document.add_picture(solicitacao.secretaria.prefeitura.get_logo())
        header_table = header.add_table(1, 2, Inches(6))
        header_table_cells = header_table.rows[0].cells
        header_table_0 = header_table_cells[0].add_paragraph()
        kh = header_table_0.add_run()
        kh.add_picture(
            solicitacao.secretaria.prefeitura.get_logo(), width=Inches(1))
        header_table_1 = header_table_cells[1].add_paragraph(
            solicitacao.secretaria.prefeitura.get_prefeitura()+'\n'+solicitacao.secretaria.nome+'\n'+solicitacao.secretaria.endereco)
        header_table_1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # p2 = header.paragraphs[1]
        # p2.text = solicitacao.secretaria.nome
        tempdir = tempfile.mkdtemp()
        document.save(os.path.join(tempdir, 'report.docx'))
        with open(os.path.join(tempdir, 'report.docx'), "rb") as doc_ok:
            response: HttpResponse = HttpResponse(doc_ok,
                                                  content_type='application/docx')
            response['Content-Disposition'] = 'attachment; filename="report.docx"'
        return response


class SolicitacaoCreateView(LoginRequiredMixin, View):
    model = Solicitacao
    template_name = 'solicitacao/form.html'
    fields = '__all__'
    success_url = 'produto:busca'

    def get(self, *args, **kwargs):
        form = SolicitacaoForm()
        return render(self.request, 'solicitacao/form.html', {'form': form})

    def post(self, *args, **kwargs):
        if self.request.user.is_authenticated:
            user = self.request.user
            data = self.request.POST

            # print(user)
            # self.request.POST['usuario'] = user
            #   'data': ['02/04/2022'], 'status': ['P']}
            dados = {
                'csrfmiddlewaretoken': data['csrfmiddlewaretoken'],
                'objeto': data['objeto'],
                'secretaria': data['secretaria'],
                'data': data['data'],
                'status': data['status'],
            }
            # pprint(dir(user))
            # pprint(vars(user))
        form = SolicitacaoForm(data=self.request.POST,
                               files=self.request.FILES)
        # pprint(dir(form.clean()))
        # pprint(vars(form.clean()))
        secretaria = Secretaria.objects.filter(pk=data['secretaria']).first()
        if form.is_valid():
            solicitacao = Solicitacao(
                usuario=user,
                objeto=data['objeto'],
                secretaria=secretaria,
                data=datetime.strptime(data['data'], '%d/%m/%Y'),
                status=data['status']
            )
            solicitacao.save()
            # form.save(commit=False)
            # form.user = user.User
            # pprint(dir(form))
            # pprint(vars(form))
            # form.save()
            return redirect(reverse('produto:busca', kwargs={'pk': solicitacao.pk}))
        else:
            return render(self.request, 'solicitacao/form.html', {'form': form})


class SolicitacaoEdit(LoginRequiredMixin, UpdateView):
    model = Solicitacao
    template_name = 'solicitacao/edit.html'
    # fields = '__all__'
    form_class = SolicitacaoForm
    # success_url = redirect('produto:busca' self)

    # def post(self, request: HttpRequest, *args: Any, **kwargs: Any) -> HttpResponse:
    #     #  redirect('produto:busca', kwargs={'pk': kwargs.get('pk')})
    #     return super().post(request, *args, **kwargs)

    def get_success_url(self, *args, **kwargs) -> str:
        x = self.kwargs.get('pk')
        return reverse('produto:busca', args=[x])


class SolicitacaoDeleteView(LoginRequiredMixin, DeleteView):
    model = Solicitacao
    template_name = "produto/delete.html"
    success_url = reverse_lazy('produto:index')


def password_reset_request(request):
    if request.method == "POST":
        password_reset_form = PasswordResetForm(request.POST)
        if password_reset_form.is_valid():
            data = password_reset_form.cleaned_data['email']
            associated_users = User.objects.filter(Q(email=data))
            if associated_users.exists():
                for user in associated_users:
                    subject = "Requisição para resetar senha - Solicitações App"
                    email_template_name = "password/templates/password_reset_email.txt"
                    c = {
                        "email": user.email,
                        # 'domain': '127.0.0.1:8000',
                        'domain': config('DOMAIN_URL'),
                        'site_name': 'Solicitação v1.0',
                        "uid": urlsafe_base64_encode(force_bytes(user.pk)),
                        "user": user,
                        'token': default_token_generator.make_token(user),
                        'protocol': config('DOMAIN_PROTOCOL'),
                    }
                    email = render_to_string(email_template_name, c)
                    try:
                        send_mail(subject, email, 'icarodavi@gmail.com',
                                  [user.email], fail_silently=False)
                    except BadHeaderError:
                        return HttpResponse('Invalid header found.')
                    return redirect("/password_reset/done/")
    password_reset_form = PasswordResetForm()
    return render(request=request, template_name="password/templates/password_reset.html", context={"password_reset_form": password_reset_form})
